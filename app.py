from flask import Flask, render_template, request, jsonify, send_file
from datetime import date, datetime
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

app = Flask(__name__)

# Global variables
TEMPLATE_PATH = "Template.docx"
OUTPUT_DIR = "Exported"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# Thai date utilities
THAI_MONTHS = [
    "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน",
    "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม",
    "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
]

def convert_to_buddhist_year(year):
    return year + 543

def format_thai_date(date_obj):
    return f"{date_obj.day} {THAI_MONTHS[date_obj.month - 1]} {convert_to_buddhist_year(date_obj.year)}"

# Department configurations
departments = {
    'กลุ่มสาระการเรียนรู้วิทยาศาสตร์และเทคโนโลยี': 'กลุ่มสาระการเรียนรู้วิทยาศาสตร์และเทคโนโลยี',
    'กลุ่มสาระการเรียนรู้ภาษาไทย': 'กลุ่มสาระการเรียนรู้ภาษาไทย',
    'กลุ่มสาระการเรียนรู้คณิตศาสตร์': 'กลุ่มสาระการเรียนรู้คณิตศาสตร์',
    'กลุ่มสาระการเรียนรู้สังคมศึกษาฯ': 'กลุ่มสาระการเรียนรู้สังคมศึกษาฯ',
    'กลุ่มสาระการเรียนรู้ภาษาต่างประเทศ': 'กลุ่มสาระการเรียนรู้ภาษาต่างประเทศ',
    'กลุ่มสาระการเรียนรู้การงานอาชีพ': 'กลุ่มสาระการเรียนรู้การงานอาชีพ',
    'กลุ่มสาระการเรียนรู้สุขศึกษาฯ': 'กลุ่มสาระการเรียนรู้สุขศึกษาฯ'
}

# Document processing functions
def set_document_font(doc, font_name="TH SarabunPSK", font_size=16):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        run.font.size = Pt(font_size)

def replace_text_in_document(doc, replacements):
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
                if key in ["{issuer_name}", "{issuer_position}"]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        for paragraph in cell.paragraphs:
                            paragraph.text = paragraph.text.replace(key, str(value))
                            if key in ["{issuer_name}", "{issuer_position}"]:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def generate_document(form_data):
    doc = Document(TEMPLATE_PATH)
    
    # Format dates
    submission_date = datetime.strptime(form_data['date'], '%Y-%m-%d').date()
    formatted_submission_date = format_thai_date(submission_date)
    
    # Format activity dates
    start_date = datetime.strptime(form_data['activity']['startDate'], '%Y-%m-%d').date()
    end_date = datetime.strptime(form_data['activity']['endDate'], '%Y-%m-%d').date()
    
    if start_date == end_date:
        activity_date_str = f"ในวันที่ {format_thai_date(start_date)}"
    else:
        activity_date_str = f"ระหว่างวันที่ {start_date.day} - {end_date.day} {THAI_MONTHS[end_date.month - 1]} พ.ศ. {convert_to_buddhist_year(end_date.year)}"
    
    # Generate participants list
    participants_list = []
    for idx, student in enumerate(form_data['students'], 1):
        participants_list.append(
            f"   {idx}. {student['title']}{student['firstname']} "
            f"{student['lastname']} นักเรียนชั้นมัธยมศึกษาปีที่ {student['grade']}/{student['room']}"
        )
    
    for idx, teacher in enumerate(form_data['teachers'], len(participants_list) + 1):
        participants_list.append(
            f"   {idx}. {teacher['title']}{teacher['firstname']} "
            f"{teacher['lastname']} {teacher['department']}"
        )
    
    combined_list = "\n".join(participants_list)
    
    # Define replacements
    replacements = {
        "{departments[selected_department]}": form_data['department'],
        "{formatted_submission_date}": formatted_submission_date,
        "{subject}": form_data['subject'],
        "{activity_name}": form_data['activity']['name'],
        "{location}": form_data['activity']['location'],
        "{activity_date_str}": activity_date_str,
        "{stdname}": combined_list,
        "{issuer_name}": form_data['issuer']['name'],
        "{issuer_position}": form_data['issuer']['position']
    }
    
    # Process document
    replace_text_in_document(doc, replacements)
    set_document_font(doc)
    
    # Save document
    output_path = os.path.join(OUTPUT_DIR, "Output.docx")
    doc.save(output_path)
    return output_path

@app.route('/generate', methods=['POST'])
def generate():
    try:
        form_data = request.get_json()
        output_path = generate_document(form_data)
        
        # Convert to PDF if requested
        if form_data.get('generate_pdf', False):
            # PDF conversion using docx2pdf requires Microsoft Word (Windows or macOS).
            # On Linux (e.g., Railway) PDF conversion via COM is not available. Return
            # a helpful message instead of attempting an unsupported operation.
            if sys.platform == 'win32' or sys.platform == 'darwin':
                from docx2pdf import convert
                import pythoncom
                pdf_path = os.path.join(OUTPUT_DIR, "Output.pdf")
                pythoncom.CoInitialize()
                try:
                    convert(output_path, pdf_path)
                finally:
                    pythoncom.CoUninitialize()
                return jsonify({"docx_path": output_path, "pdf_path": pdf_path})
            else:
                # Non-Windows/macOS runtime: skip PDF conversion and inform caller.
                return jsonify({
                    "docx_path": output_path,
                    "pdf_path": None,
                    "warning": "PDF conversion is not available on this platform."
                })
        
        return jsonify({"docx_path": output_path})
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/download/<filetype>')
def download(filetype):
    if filetype not in ['docx', 'pdf']:
        return "Invalid file type", 400
    
    filename = f"Output.{filetype}"
    filepath = os.path.join(OUTPUT_DIR, filename)
    
    if not os.path.exists(filepath):
        return "File not found", 404
    
    mime_types = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf'
    }
    
    return send_file(filepath,
                    mimetype=mime_types[filetype],
                    as_attachment=True,
                    download_name=filename)


@app.route('/')
def landing():
    return render_template('landing.html', year=datetime.now().year)

@app.route('/memo')
def memo():
    today = date.today()
    return render_template('form1.html',
                         title="Webapp สร้างบันทึกข้อความ",
                         departments=departments,
                         today=today)

if __name__ == '__main__':
    app.run(debug=True)